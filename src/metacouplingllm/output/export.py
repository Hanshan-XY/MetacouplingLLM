"""Scholar-facing Markdown + Word exports for ``AnalysisResult``.

PR #31: turns the existing ``AnalysisResult`` (rich Python object with
``parsed``, ``formatted`` text, optional ``map`` Figure) into formats
scholars can drop straight into a paper draft:

- ``render_markdown(result, path=None)`` -- single Markdown string.
  When ``path`` is given, also writes the .md file and saves the map
  as ``<stem>_map.png`` alongside.
- ``render_docx(result, path=None)`` -- Word .docx file with headings,
  flow / evidence tables, and the map embedded.  Requires the optional
  ``python-docx`` dependency (``pip install metacouplingllm[export]``).

Both renderers read from the SAME structured intermediate
(``_build_sections``) so the two formats stay in sync.  Sections is
internal; not exposed on ``AnalysisResult`` to keep the public API
small per the PR #31 design discussion ("user-friendly way to show
the output, most extra work is not necessary").
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:  # pragma: no cover - type-only
    from metacouplingllm.core import AnalysisResult


# ---------------------------------------------------------------------------
# Internal intermediate -- structured sections built from ParsedAnalysis
# ---------------------------------------------------------------------------


def _build_sections(result: "AnalysisResult") -> dict[str, Any]:
    """Pull structured pieces of the analysis into one dict for both
    renderers.  Reads ``result.parsed`` (ParsedAnalysis) directly rather
    than re-parsing ``result.formatted`` -- the data is already
    structured.

    Returns
    -------
    Dict shape:

    ```
    {
        "abstract": str,
        "coupling_classification": str,
        "intracoupling": dict | None,    # CouplingSection as dict
        "pericoupling": dict | None,
        "telecoupling": dict | None,
        "cross_coupling_interactions": list[str],
        "research_gaps": list[str],
        "evidence_coverage": str,
        "flows": list[dict],             # from map_data or fallback
        "web_sources": list[dict],       # [{id, title, url, model_summary}]
        "focal_country": str | None,
        "topic": str,                    # best-effort topic label
    }
    ```
    """
    parsed = result.parsed

    def _section_to_dict(section: Any) -> dict[str, Any] | None:
        if section is None or getattr(section, "is_empty", True):
            return None
        return {
            "systems": list(getattr(section, "systems", []) or []),
            "flows": list(getattr(section, "flows", []) or []),
            "agents": list(getattr(section, "agents", []) or []),
            "causes": dict(getattr(section, "causes", {}) or {}),
            "effects": dict(getattr(section, "effects", {}) or {}),
        }

    # Flows table: prefer Stage-3 structured map_data; fall back to
    # iterating coupling-section flow prose entries.
    flows: list[dict[str, Any]] = []
    map_data = getattr(parsed, "map_data", None)
    if isinstance(map_data, dict):
        raw_flows = map_data.get("flows") or []
        if isinstance(raw_flows, list):
            for flow in raw_flows:
                if not isinstance(flow, dict):
                    continue
                flows.append({
                    "category": str(flow.get("category", "")).strip(),
                    "source": str(flow.get("source", "")).strip(),
                    "target": str(flow.get("target", "")).strip(),
                    "bidirectional": bool(flow.get("bidirectional", False)),
                    "description": str(flow.get("description", "")).strip(),
                })
    if not flows:
        # Fallback: walk the parsed CouplingSection prose flows.  These
        # don't have ISO codes, just direction strings, but at least
        # something renders in the table.
        for section_name in ("intracoupling", "pericoupling", "telecoupling"):
            section = getattr(parsed, section_name, None)
            if section is None:
                continue
            for flow in getattr(section, "flows", []) or []:
                if not isinstance(flow, dict):
                    continue
                flows.append({
                    "category": str(flow.get("category", "")).strip(),
                    "source": "",
                    "target": "",
                    "bidirectional": False,
                    "description": str(flow.get("direction", "")).strip()
                    + (
                        (" — " + str(flow.get("description", "")).strip())
                        if flow.get("description")
                        else ""
                    ),
                })

    # Web sources: pull from result.web_map_signals or the assistant's
    # _last_web_results if surfaced.  Use the raw evidence list when
    # available; the IDs (W1, W2, ...) match the markers in formatted.
    web_sources: list[dict[str, str]] = []
    raw_web = None
    if isinstance(result.web_map_signals, dict):
        raw_web = result.web_map_signals.get("evidence")
    # Some assistants expose the raw list on the result via a
    # separate attribute populated by the pipeline; check defensively.
    if not raw_web:
        raw_web = getattr(result, "_web_sources_for_export", None)
    if isinstance(raw_web, list):
        for idx, src in enumerate(raw_web, 1):
            if not isinstance(src, dict):
                continue
            web_sources.append({
                "id": f"W{idx}",
                "title": str(src.get("title", "")).strip(),
                "url": str(src.get("url", "")).strip(),
                "model_summary": str(
                    src.get("model_summary", src.get("snippet", ""))
                ).strip(),
            })

    focal_country: str | None = None
    if isinstance(map_data, dict):
        fc = map_data.get("focal_country")
        if isinstance(fc, str) and fc.strip():
            focal_country = fc.strip()

    # PR #31: RAG retrieval hits.  Populated by ``_build_result()``
    # when the pipeline ran with an RAG engine in pre_retrieval mode.
    # Each entry becomes a row in the "Evidence from Literature"
    # section.  Empty when the run had no RAG engine or post_hoc
    # mode (which doesn't attach hits to the result).
    rag_hits: list[dict[str, Any]] = []
    raw_rag = getattr(result, "_rag_hits_for_export", None)
    if isinstance(raw_rag, list):
        for idx, hit in enumerate(raw_rag, 1):
            chunk = getattr(hit, "chunk", None)
            if chunk is None:
                continue
            # Use the existing display-path excerpt cap (~300 chars)
            # so the docx / markdown shows what scholars see in the
            # text formatter, not the full chunk.
            text = str(getattr(chunk, "text", "")).strip()
            if len(text) > 600:
                text = text[:600].rstrip() + "..."
            rag_hits.append({
                "id": f"T{result.turn_number}:{idx}",
                "paper_title": str(getattr(chunk, "paper_title", "")).strip(),
                "authors": str(getattr(chunk, "authors", "")).strip(),
                "year": str(getattr(chunk, "year", "") or "").strip(),
                "section": str(getattr(chunk, "section", "")).strip(),
                "score": float(getattr(hit, "score", 0.0) or 0.0),
                "text": text,
            })

    # PR #32: prefer the user's original query for the title.  Falls
    # back to the coupling_classification-derived "topic" heuristic
    # only for results constructed outside the analyze() pipeline
    # (e.g., direct AnalysisResult(...) calls in unit tests).
    original_query = (
        getattr(result, "_original_query_for_export", "") or ""
    ).strip()

    # Topic heuristic (fallback): first 80 chars of
    # coupling_classification, or a generic label.
    topic = ""
    if parsed.coupling_classification:
        first = parsed.coupling_classification.split("\n")[0].strip()
        topic = first[:80] + ("…" if len(first) > 80 else "")
    if not topic:
        topic = "metacoupling analysis"

    return {
        "abstract": getattr(result, "abstract", "") or "",
        "coupling_classification": parsed.coupling_classification or "",
        "intracoupling": _section_to_dict(
            getattr(parsed, "intracoupling", None)
        ),
        "pericoupling": _section_to_dict(
            getattr(parsed, "pericoupling", None)
        ),
        "telecoupling": _section_to_dict(
            getattr(parsed, "telecoupling", None)
        ),
        "cross_coupling_interactions": list(
            getattr(parsed, "cross_coupling_interactions", []) or []
        ),
        "research_gaps": list(
            getattr(parsed, "research_gaps", []) or []
        ),
        "evidence_coverage": (
            getattr(parsed, "evidence_coverage_note", "") or ""
        ),
        "flows": flows,
        "web_sources": web_sources,
        "rag_hits": rag_hits,
        "focal_country": focal_country,
        "topic": topic,
        "original_query": original_query,
    }


# ---------------------------------------------------------------------------
# Helpers shared by both renderers
# ---------------------------------------------------------------------------

# PR #32: canonical Liu 2017 flow-category order, used to group §N.2
# Flows into per-category subsections (§N.2.1 Matter, §N.2.2 Capital,
# etc.).  Mirrors ``_CANONICAL_FLOW_CATEGORIES`` in core.py but
# preserves the canonical *order* (frozenset there is unordered).
# Unknown categories the LLM emits get appended last under their
# original label.
_CANONICAL_FLOW_CATEGORY_ORDER: tuple[str, ...] = (
    "matter",
    "capital",
    "information",
    "energy",
    "people",
    "organisms",
)


def _group_flows_by_category(
    flows: list[dict[str, Any]],
) -> list[tuple[str, list[dict[str, Any]]]]:
    """Group a flat flows list by category, preserving canonical
    Liu 2017 order.  Unknown categories appended last in
    insertion order.  Returns ``[(display_category, flows_in_group),
    ...]`` skipping empty groups.

    The display category is title-cased (``"matter"`` → ``"Matter"``)
    to match the existing AnalysisFormatter convention.
    """
    by_cat: dict[str, list[dict[str, Any]]] = {}
    unknown_order: list[str] = []  # preserve LLM emission order
    for flow in flows:
        raw = str(flow.get("category", "")).strip().lower()
        if not raw:
            raw = "unspecified"
        if raw not in by_cat:
            if raw not in _CANONICAL_FLOW_CATEGORY_ORDER:
                unknown_order.append(raw)
            by_cat[raw] = []
        by_cat[raw].append(flow)

    grouped: list[tuple[str, list[dict[str, Any]]]] = []
    for cat in _CANONICAL_FLOW_CATEGORY_ORDER:
        if cat in by_cat and by_cat[cat]:
            grouped.append((cat.title(), by_cat[cat]))
    for cat in unknown_order:
        if by_cat[cat]:
            grouped.append((cat.title(), by_cat[cat]))
    return grouped


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------


def _md_coupling_section(
    parts: list[str], number: int, title: str, section: dict[str, Any] | None,
) -> None:
    """Render one coupling section as Markdown subsections.

    PR #32 polish changes applied here:
    - §N.1 Systems sub-field labels (Human subsystem, etc.) are
      bold (was italic).
    - §N.2 Flows grouped by canonical Liu 2017 category into
      §N.2.K subsections instead of one flat numbered list.
    - §N.3 Agents level prefix bold (was italic).
    - §N.4 Causes / §N.5 Effects categories rendered as
      §N.4.K / §N.5.K subsections instead of inline bullet blocks.
    """
    if section is None:
        return
    parts.append(f"## {number}. {title}")
    parts.append("")

    if section["systems"]:
        parts.append(f"### {number}.1 Systems")
        for entry in section["systems"]:
            role = str(entry.get("role", "system")).title()
            name = str(entry.get("name", "")).strip()
            parts.append(f"- **{role}**" + (f": {name}" if name else ""))
            for key in (
                "human_subsystem", "natural_subsystem",
                "geographic_scope", "description",
            ):
                val = str(entry.get(key, "")).strip()
                if val:
                    label = key.replace("_", " ").capitalize()
                    # PR #32: bold (was italic).
                    parts.append(f"  - **{label}**: {val}")
        parts.append("")

    if section["flows"]:
        parts.append(f"### {number}.2 Flows")
        parts.append("")
        # PR #32: group by category with §N.2.K subsections.
        groups = _group_flows_by_category(section["flows"])
        for sub_idx, (category, group_flows) in enumerate(groups, 1):
            parts.append(f"#### {number}.2.{sub_idx} {category}")
            for idx, flow in enumerate(group_flows, 1):
                direction = str(flow.get("direction", "")).strip()
                description = str(flow.get("description", "")).strip()
                line = f"{idx}."
                if direction:
                    line += f" {direction}"
                if description:
                    if direction:
                        line += f" — {description}"
                    else:
                        line += f" {description}"
                if not direction and not description:
                    line += " (no details)"
                parts.append(line)
            parts.append("")

    if section["agents"]:
        parts.append(f"### {number}.3 Agents")
        for agent in section["agents"]:
            level = str(agent.get("level", "")).title()
            name = str(agent.get("name", "")).strip()
            desc = str(agent.get("description", "")).strip()
            # PR #32: bold (was italic).
            prefix = f"**{level}** " if level else ""
            suffix = f" — {desc}" if desc else ""
            parts.append(f"- {prefix}{name}{suffix}")
        parts.append("")

    # PR #32: per-category §N.4.K / §N.5.K subsections instead of
    # inline `- **Category**:` bullets.  Matches the Flows treatment
    # and visually separates each category in both Markdown and docx.
    for sub_idx, key, label in (
        (4, "causes", "Causes"),
        (5, "effects", "Effects"),
    ):
        items = section[key]
        if not items:
            continue
        parts.append(f"### {number}.{sub_idx} {label}")
        parts.append("")
        for cat_idx, (category, entries) in enumerate(items.items(), 1):
            parts.append(
                f"#### {number}.{sub_idx}.{cat_idx} {category.title()}"
            )
            for entry in entries:
                parts.append(f"- {entry}")
            parts.append("")


def render_markdown(
    result: "AnalysisResult", path: str | Path | None = None,
) -> str:
    """Convert ``result`` into a Markdown string.  Writes a file +
    saves the map as a sibling PNG when ``path`` is given."""
    s = _build_sections(result)
    parts: list[str] = []

    # PR #32: title is "Metacoupling Analysis: {user query}".  Falls
    # back to the older "{focal}: {topic}" heuristic only when no
    # original_query is attached (e.g., direct AnalysisResult(...)
    # construction in tests).
    if s["original_query"]:
        parts.append(f"# Metacoupling Analysis: {s['original_query']}")
    else:
        focal = s["focal_country"] or "Case"
        parts.append(f"# Metacoupling Analysis — {focal}: {s['topic']}")
    parts.append("")

    # Abstract
    if s["abstract"]:
        parts.append("## Abstract")
        parts.append("")
        parts.append(s["abstract"])
        parts.append("")

    # §1 Classification
    if s["coupling_classification"]:
        parts.append("## 1. Coupling Classification")
        parts.append("")
        parts.append(s["coupling_classification"])
        parts.append("")

    # §2-§4 Coupling sections
    _md_coupling_section(parts, 2, "Intracoupling Analysis", s["intracoupling"])
    _md_coupling_section(parts, 3, "Pericoupling Analysis", s["pericoupling"])
    _md_coupling_section(parts, 4, "Telecoupling Analysis", s["telecoupling"])

    # §5 Cross-coupling
    if s["cross_coupling_interactions"]:
        parts.append("## 5. Cross-Coupling Interactions")
        parts.append("")
        for item in s["cross_coupling_interactions"]:
            parts.append(f"- {item}")
        parts.append("")

    # §6 Research gaps
    if s["research_gaps"]:
        parts.append("## 6. Research Gaps")
        parts.append("")
        for gap in s["research_gaps"]:
            parts.append(f"- {gap}")
        parts.append("")

    # §7 Evidence coverage
    if s["evidence_coverage"]:
        parts.append("## 7. Evidence Coverage")
        parts.append("")
        parts.append(s["evidence_coverage"])
        parts.append("")

    # Flows table
    if s["flows"]:
        parts.append("## Flows")
        parts.append("")
        parts.append("| Category | Source | Target | Bidirectional | Description |")
        parts.append("|---|---|---|---|---|")
        for flow in s["flows"]:
            # Title-case the category to match the existing
            # AnalysisFormatter convention (so "matter" → "Matter"
            # consistent with the text formatter PR #25 / PR #26 use).
            cat = (flow["category"] or "—").title()
            src = flow["source"] or "—"
            tgt = flow["target"] or "—"
            bi = "Yes" if flow["bidirectional"] else "No"
            desc = flow["description"] or ""
            # Escape pipes inside cells to keep the table valid.
            desc = desc.replace("|", "\\|")
            parts.append(f"| {cat} | {src} | {tgt} | {bi} | {desc} |")
        parts.append("")

    # Evidence from Literature (RAG passages, when present)
    if s["rag_hits"]:
        parts.append("## Evidence from Literature")
        parts.append("")
        for hit in s["rag_hits"]:
            title = hit["paper_title"] or "(untitled paper)"
            citation_bits: list[str] = []
            if hit["authors"]:
                citation_bits.append(hit["authors"])
            if hit["year"]:
                citation_bits.append(str(hit["year"]))
            citation = ", ".join(citation_bits)
            parts.append(f"### [{hit['id']}] {title}")
            if citation:
                parts.append(f"*{citation}*")
            if hit["section"]:
                parts.append(f"Section: {hit['section']}")
            if hit["text"]:
                parts.append("")
                # Use blockquote so excerpts stand out from prose.
                for line in hit["text"].splitlines():
                    parts.append(f"> {line}")
            parts.append("")

    # Web sources
    if s["web_sources"]:
        parts.append("## Web Sources")
        parts.append("")
        for src in s["web_sources"]:
            line = f"- **[{src['id']}]** {src['title']}"
            if src["url"]:
                line += f" — <{src['url']}>"
            parts.append(line)
        parts.append("")

    # Map embed (only when we have a file path to write the PNG to).
    map_rel_name: str | None = None
    if path is not None and result.map is not None:
        path = Path(path)
        map_rel_name = f"{path.stem}_map.png"
        try:
            result.map.savefig(
                path.parent / map_rel_name,
                dpi=150,
                bbox_inches="tight",
            )
            parts.append("## Map")
            parts.append("")
            parts.append(f"![map]({map_rel_name})")
            parts.append("")
        except Exception as exc:  # pragma: no cover - matplotlib edge case
            parts.append(f"<!-- Map render failed: {exc} -->")
            parts.append("")

    text = "\n".join(parts).rstrip() + "\n"

    if path is not None:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding="utf-8")

    return text


# ---------------------------------------------------------------------------
# Word (.docx) renderer -- requires python-docx
# ---------------------------------------------------------------------------


def render_docx(
    result: "AnalysisResult", path: str | Path | None = None,
) -> Path:
    """Convert ``result`` into a Word document.  Returns the
    ``pathlib.Path`` written.  Raises ``ImportError`` when
    ``python-docx`` is not installed."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as exc:  # pragma: no cover - exercised in tests
        raise ImportError(
            "render_docx requires the optional ``python-docx`` "
            "dependency.  Install it with "
            "``pip install metacouplingllm[export]``."
        ) from exc

    s = _build_sections(result)
    out_path = Path(path) if path is not None else Path("metacoupling_report.docx")

    doc = Document()

    # PR #32: title is "Metacoupling Analysis: {user query}" (see
    # _build_sections for fallback).
    if s["original_query"]:
        title_text = f"Metacoupling Analysis: {s['original_query']}"
    else:
        focal = s["focal_country"] or "Case"
        title_text = f"Metacoupling Analysis — {focal}: {s['topic']}"
    doc.add_heading(title_text, level=0)

    # Abstract
    if s["abstract"]:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(s["abstract"])

    # §1 Classification
    if s["coupling_classification"]:
        doc.add_heading("1. Coupling Classification", level=1)
        doc.add_paragraph(s["coupling_classification"])

    # §2-§4 Coupling sections
    for number, title, section in (
        (2, "Intracoupling Analysis", s["intracoupling"]),
        (3, "Pericoupling Analysis", s["pericoupling"]),
        (4, "Telecoupling Analysis", s["telecoupling"]),
    ):
        if section is None:
            continue
        doc.add_heading(f"{number}. {title}", level=1)
        if section["systems"]:
            doc.add_heading(f"{number}.1 Systems", level=2)
            for entry in section["systems"]:
                role = str(entry.get("role", "system")).title()
                name = str(entry.get("name", "")).strip()
                p = doc.add_paragraph()
                run = p.add_run(role)
                run.bold = True
                if name:
                    p.add_run(f": {name}")
                for key in (
                    "human_subsystem", "natural_subsystem",
                    "geographic_scope", "description",
                ):
                    val = str(entry.get(key, "")).strip()
                    if val:
                        label = key.replace("_", " ").capitalize()
                        # PR #32: bold the sub-field label (was a
                        # plain non-bold paragraph) so scholars can
                        # scan the document fast.
                        bp = doc.add_paragraph(style="List Bullet")
                        brun = bp.add_run(f"{label}: ")
                        brun.bold = True
                        bp.add_run(val)
        if section["flows"]:
            doc.add_heading(f"{number}.2 Flows", level=2)
            # PR #32: per-category §N.2.K subsections instead of one
            # flat numbered list.  Avoids the visual mess of 15+
            # mixed-category flows.
            groups = _group_flows_by_category(section["flows"])
            for sub_idx, (category, group_flows) in enumerate(groups, 1):
                doc.add_heading(
                    f"{number}.2.{sub_idx} {category}", level=3,
                )
                for flow in group_flows:
                    direction = str(flow.get("direction", "")).strip()
                    description = str(flow.get("description", "")).strip()
                    text = ""
                    if direction:
                        text += direction
                    if description:
                        text += (f" — {description}" if direction
                                 else description)
                    if not text:
                        text = "(no details)"
                    doc.add_paragraph(text, style="List Number")
        if section["agents"]:
            doc.add_heading(f"{number}.3 Agents", level=2)
            for agent in section["agents"]:
                level = str(agent.get("level", "")).title()
                name = str(agent.get("name", "")).strip()
                desc = str(agent.get("description", "")).strip()
                # PR #32: bold the level label (e.g.,
                # "Individuals / Households", "Firms / Traders /
                # Corporations") so scholars can scan the document
                # by agent class.
                p = doc.add_paragraph(style="List Bullet")
                if level:
                    lrun = p.add_run(f"{level} ")
                    lrun.bold = True
                p.add_run(name)
                if desc:
                    p.add_run(f" — {desc}")
        # PR #32: per-category §N.4.K / §N.5.K subsections (Heading
        # 3) instead of one paragraph per category with semicolon-
        # joined items.  Each item becomes its own bullet.
        for sub_idx, key, label in (
            (4, "causes", "Causes"),
            (5, "effects", "Effects"),
        ):
            items = section[key]
            if not items:
                continue
            doc.add_heading(f"{number}.{sub_idx} {label}", level=2)
            for cat_idx, (category, entries) in enumerate(items.items(), 1):
                doc.add_heading(
                    f"{number}.{sub_idx}.{cat_idx} {category.title()}",
                    level=3,
                )
                for entry in entries:
                    doc.add_paragraph(entry, style="List Bullet")

    # §5 Cross-coupling
    if s["cross_coupling_interactions"]:
        doc.add_heading("5. Cross-Coupling Interactions", level=1)
        for item in s["cross_coupling_interactions"]:
            doc.add_paragraph(item, style="List Bullet")

    # §6 Research gaps
    if s["research_gaps"]:
        doc.add_heading("6. Research Gaps", level=1)
        for gap in s["research_gaps"]:
            doc.add_paragraph(gap, style="List Bullet")

    # §7 Evidence coverage
    if s["evidence_coverage"]:
        doc.add_heading("7. Evidence Coverage", level=1)
        doc.add_paragraph(s["evidence_coverage"])

    # Flows table
    if s["flows"]:
        doc.add_heading("Flows", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Source"
        hdr[2].text = "Target"
        hdr[3].text = "Bidirectional"
        hdr[4].text = "Description"
        for flow in s["flows"]:
            row = table.add_row().cells
            row[0].text = (flow["category"] or "—").title()
            row[1].text = flow["source"] or "—"
            row[2].text = flow["target"] or "—"
            row[3].text = "Yes" if flow["bidirectional"] else "No"
            row[4].text = flow["description"] or ""

    # Evidence from Literature (RAG passages, when present).  Rendered
    # before the Web Sources table so the document flows
    # literature → web → map, matching scholar expectations.
    if s["rag_hits"]:
        doc.add_heading("Evidence from Literature", level=1)
        for hit in s["rag_hits"]:
            title = hit["paper_title"] or "(untitled paper)"
            doc.add_heading(f"[{hit['id']}] {title}", level=2)
            # Author / year line in italics.
            citation_bits: list[str] = []
            if hit["authors"]:
                citation_bits.append(hit["authors"])
            if hit["year"]:
                citation_bits.append(str(hit["year"]))
            if citation_bits:
                p = doc.add_paragraph()
                run = p.add_run(", ".join(citation_bits))
                run.italic = True
            if hit["section"]:
                doc.add_paragraph(f"Section: {hit['section']}")
            if hit["text"]:
                # Use the Quote style so the excerpt visually
                # separates from prose; falls back to a regular
                # paragraph on templates without Quote.
                try:
                    doc.add_paragraph(hit["text"], style="Quote")
                except KeyError:
                    doc.add_paragraph(hit["text"])

    # Web sources table
    if s["web_sources"]:
        doc.add_heading("Web Sources", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Title"
        hdr[2].text = "URL"
        for src in s["web_sources"]:
            row = table.add_row().cells
            row[0].text = src["id"]
            row[1].text = src["title"]
            row[2].text = src["url"]

    # Map figure
    if result.map is not None:
        doc.add_heading("Map", level=1)
        buf = BytesIO()
        try:
            result.map.savefig(buf, format="png", dpi=150, bbox_inches="tight")
            buf.seek(0)
            doc.add_picture(buf, width=Inches(6.0))
        except Exception as exc:  # pragma: no cover - matplotlib edge case
            doc.add_paragraph(f"[Map render failed: {exc}]")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
