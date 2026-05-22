"""Tests for PR #31 scholar-facing export: render_markdown / render_docx
and the AnalysisResult.to_markdown() / .to_docx() convenience methods.

Built against ``output/export.py``.  ``python-docx`` is optional --
docx tests skip when the dependency isn't installed.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from metacouplingllm.core import AnalysisResult
from metacouplingllm.llm.parser import ParsedAnalysis
from metacouplingllm.output.export import (
    _build_sections,
    _merge_fragmented_flows,
    _split_collapsed_causes_effects,
    render_docx,
    render_markdown,
)

from ._helpers import make_parsed_analysis


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _fake_rag_hit(
    title: str, authors: str, year: int, section: str, text: str, score: float,
):
    """Build a fake RetrievalResult-shaped object for tests without
    importing the real class (keeps test isolated from RAG internals)."""
    from types import SimpleNamespace

    chunk = SimpleNamespace(
        paper_key=title.lower().replace(" ", "_"),
        paper_title=title,
        authors=authors,
        year=year,
        section=section,
        text=text,
        chunk_index=0,
    )
    return SimpleNamespace(chunk=chunk, score=score)


@pytest.fixture
def minimal_result() -> AnalysisResult:
    """An AnalysisResult with just enough content for the renderers."""
    parsed = make_parsed_analysis(
        coupling_classification="Telecoupling (Mexico → USA avocado trade).",
        systems={
            "sending": {
                "name": "Mexico (Jalisco)",
                "geographic_scope": "Western Mexico",
                "description": "Major avocado-producing region.",
            },
            "receiving": {
                "name": "United States",
                "geographic_scope": "Continental USA",
                "description": "Primary avocado importer.",
            },
        },
        flows=[
            {
                "category": "matter",
                "direction": "Mexico → USA",
                "description": "Fresh avocado exports",
            },
            {
                "category": "capital",
                "direction": "USA → Mexico",
                "description": "Payments for produce",
            },
            {
                "category": "information",
                "direction": "USDA → SENASICA",
                "description": "Phytosanitary protocols",
            },
        ],
        agents=[
            {
                "level": "Individuals / Households",
                "name": "Smallholder growers",
                "description": "Manage orchards.",
            },
            {
                "level": "Firms / Traders / Corporations",
                "name": "Packhouses / exporters",
                "description": "Cold-chain logistics.",
            },
        ],
        causes={
            "Economic": [
                "Rising US demand for avocados.",
                "Favorable export prices.",
            ],
            "Political / Institutional": [
                "APHIS-SENASICA bilateral protocol (2022).",
            ],
        },
        effects={
            "Economic": [
                "Reinvestment in orchard expansion.",
            ],
            "Environmental": [
                "Deforestation pressure in highland forests.",
                "Increased irrigation demand.",
            ],
        },
        coupling_type="telecoupling",
        raw_text="Avocado telecoupling analysis.",
    )
    parsed.cross_coupling_interactions = [
        "Trade tariffs feed back into Mexican producer prices.",
    ]
    parsed.research_gaps = [
        "Subnational variation in Jalisco's avocado supply chain.",
    ]
    parsed.evidence_coverage_note = (
        "Coverage is strong for trade volumes; weak on farmworker welfare."
    )
    parsed.map_data = {
        "focal_country": "MEX",
        "flows": [
            {
                "category": "matter",
                "source": "MEX",
                "target": "USA",
                "bidirectional": False,
                "description": "avocado exports",
            },
            {
                "category": "capital",
                "source": "USA",
                "target": "MEX",
                "bidirectional": False,
                "description": "payments",
            },
        ],
    }

    result = AnalysisResult(
        parsed=parsed,
        formatted="(formatted-text placeholder)",
        raw="(raw-text placeholder)",
        turn_number=1,
        abstract=(
            "This study examines the Mexico-USA avocado trade as a "
            "telecoupling. Jalisco's producers send fresh fruit to U.S. "
            "consumers in exchange for capital. Key findings include "
            "tariff feedback effects on producer prices."
        ),
    )
    # PR #32: attach the user's original query so the exporters use
    # it as the document title.
    result._original_query_for_export = (
        "Mexico avocado exports to the United States, focusing on "
        "the Jalisco production region"
    )
    # Mock web sources via the assistant-side attribute used by the
    # builder.  Real pipeline populates this from _last_web_results.
    result._web_sources_for_export = [
        {
            "title": "Mexico avocado exports 2024",
            "url": "https://example.com/mex-exports-2024",
            "model_summary": "Mexico shipped 1.1M tonnes of avocados to USA.",
        },
        {
            "title": "USDA avocado import data",
            "url": "https://usda.gov/avocado-imports",
            "model_summary": "USDA records the breakdown by destination state.",
        },
    ]
    # PR #31 follow-up: RAG hits attached the same way.  Real pipeline
    # attaches result._rag_hits_for_export = list(self._last_rag_hits).
    result._rag_hits_for_export = [
        _fake_rag_hit(
            title="Mexico's avocado industry and global trade",
            authors="Garcia, M. and Hernandez, L.",
            year=2023,
            section="DISCUSSION",
            text=(
                "Jalisco's accession to the USDA Systems Approach in 2022 "
                "represents a structural shift in the bi-national avocado "
                "supply chain, with cold-chain logistics concentrated at "
                "the Laredo, TX corridor."
            ),
            score=0.82,
        ),
        _fake_rag_hit(
            title="Telecoupled agricultural systems",
            authors="Liu, J. et al.",
            year=2013,
            section="INTRODUCTION",
            text=(
                "Telecoupling links distant systems through flows of "
                "matter, capital, information, energy, people, and "
                "organisms."
            ),
            score=0.74,
        ),
    ]
    return result


# ---------------------------------------------------------------------------
# AnalysisResult.abstract field
# ---------------------------------------------------------------------------


class TestAbstractField:
    def test_abstract_default_empty_string(self):
        """PR #31: AnalysisResult.abstract defaults to empty string so
        tests / callers that don't go through _build_result don't need
        to set it."""
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="",
            raw="",
            turn_number=1,
        )
        assert result.abstract == ""

    def test_abstract_field_carries_text(self):
        text = "A focused abstract for a paper introduction."
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="",
            raw="",
            turn_number=1,
            abstract=text,
        )
        assert result.abstract == text


# ---------------------------------------------------------------------------
# _build_sections (internal intermediate)
# ---------------------------------------------------------------------------


class TestBuildSections:
    def test_pulls_abstract_from_result(self, minimal_result):
        s = _build_sections(minimal_result)
        assert "telecoupling" in s["abstract"].lower()

    def test_pulls_coupling_classification(self, minimal_result):
        s = _build_sections(minimal_result)
        assert "Telecoupling" in s["coupling_classification"]

    def test_pulls_flows_from_map_data(self, minimal_result):
        s = _build_sections(minimal_result)
        assert len(s["flows"]) == 2
        assert s["flows"][0]["category"] == "matter"
        assert s["flows"][0]["source"] == "MEX"
        assert s["flows"][0]["target"] == "USA"

    def test_pulls_web_sources_with_w_prefixed_ids(self, minimal_result):
        s = _build_sections(minimal_result)
        assert len(s["web_sources"]) == 2
        assert s["web_sources"][0]["id"] == "W1"
        assert s["web_sources"][1]["id"] == "W2"
        assert "usda.gov" in s["web_sources"][1]["url"]

    def test_focal_country_from_map_data(self, minimal_result):
        assert _build_sections(minimal_result)["focal_country"] == "MEX"

    def test_flows_falls_back_to_coupling_section_when_no_map_data(self):
        parsed = make_parsed_analysis(
            flows=[
                {
                    "category": "matter",
                    "direction": "A → B",
                    "description": "stuff",
                },
            ],
            coupling_type="telecoupling",
        )
        # No map_data attribute set.
        parsed.map_data = None
        result = AnalysisResult(
            parsed=parsed, formatted="", raw="", turn_number=1,
        )
        s = _build_sections(result)
        assert len(s["flows"]) == 1
        assert s["flows"][0]["category"] == "matter"
        # Source/target empty in fallback path.
        assert s["flows"][0]["source"] == ""

    def test_graceful_when_parsed_minimal(self):
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="", raw="", turn_number=1,
        )
        s = _build_sections(result)
        assert s["abstract"] == ""
        assert s["coupling_classification"] == ""
        assert s["intracoupling"] is None
        assert s["telecoupling"] is None
        assert s["flows"] == []
        assert s["web_sources"] == []
        assert s["rag_hits"] == []

    # PR #31 follow-up: RAG hits in the sections dict.

    def test_pulls_rag_hits_with_turn_scoped_ids(self, minimal_result):
        s = _build_sections(minimal_result)
        assert len(s["rag_hits"]) == 2
        # IDs follow the turn-scoped citation grammar [T<turn>:<idx>]
        # matching the inline citations in result.formatted.
        assert s["rag_hits"][0]["id"] == "T1:1"
        assert s["rag_hits"][1]["id"] == "T1:2"
        assert s["rag_hits"][0]["paper_title"].startswith("Mexico's avocado")
        assert s["rag_hits"][0]["year"] == "2023"
        assert s["rag_hits"][0]["score"] == pytest.approx(0.82)

    def test_rag_text_truncated_at_600_chars(self):
        long_text = "X" * 1200
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="", raw="", turn_number=1,
        )
        result._rag_hits_for_export = [
            _fake_rag_hit("T", "A", 2024, "S", long_text, 0.9),
        ]
        s = _build_sections(result)
        # Truncation cap is 600 chars + "..." suffix.
        assert len(s["rag_hits"][0]["text"]) <= 603
        assert s["rag_hits"][0]["text"].endswith("...")


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------


class TestRenderMarkdown:
    def test_returns_string_with_top_level_title(self, minimal_result):
        md = render_markdown(minimal_result)
        # Title line starts with single `# `.
        first_line = md.splitlines()[0]
        assert first_line.startswith("# Metacoupling Analysis")
        # PR #32: title contains the user's original query verbatim
        # (the fixture sets _original_query_for_export to a Jalisco
        # avocado prompt).
        assert "Mexico avocado exports" in first_line

    def test_includes_abstract_section(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## Abstract" in md
        assert "telecoupling" in md.lower()

    def test_includes_coupling_classification_section(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## 1. Coupling Classification" in md

    def test_includes_telecoupling_section(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## 4. Telecoupling Analysis" in md

    def test_includes_flows_table(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## Flows" in md
        # Table header row + at least one matter row.
        assert "| Category | Source | Target | Bidirectional |" in md
        assert "| Matter | MEX | USA |" in md

    def test_includes_web_sources_with_w_ids(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## Web Sources" in md
        assert "[W1]" in md
        assert "https://example.com/mex-exports-2024" in md

    def test_includes_research_gaps_when_present(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "## 6. Research Gaps" in md
        assert "Jalisco" in md

    def test_skips_empty_sections(self):
        # No abstract, no classification, no sections, no flows.
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="", raw="", turn_number=1,
        )
        md = render_markdown(result)
        assert "## Abstract" not in md
        assert "## 1. Coupling Classification" not in md
        assert "## Flows" not in md
        # The title is always present.
        assert md.startswith("# Metacoupling Analysis")

    def test_writes_file_when_path_given(self, minimal_result, tmp_path):
        out_path = tmp_path / "subdir" / "report.md"
        text = render_markdown(minimal_result, path=out_path)
        assert out_path.exists()
        assert out_path.read_text(encoding="utf-8") == text

    def test_returns_markdown_without_writing_when_path_none(
        self, minimal_result, tmp_path,
    ):
        # No path argument → no file should exist after the call.
        before = list(tmp_path.iterdir())
        text = render_markdown(minimal_result)
        after = list(tmp_path.iterdir())
        assert before == after
        assert "## Abstract" in text

    def test_to_markdown_method_delegates(self, minimal_result):
        md_direct = render_markdown(minimal_result)
        md_method = minimal_result.to_markdown()
        assert md_direct == md_method

    # PR #31 follow-up: Evidence from Literature rendering.

    def test_includes_rag_evidence_section_when_hits_present(
        self, minimal_result,
    ):
        md = render_markdown(minimal_result)
        assert "## Evidence from Literature" in md
        # Each hit gets its own ### heading with the turn-scoped ID.
        assert "### [T1:1]" in md
        assert "### [T1:2]" in md
        # Paper title visible.
        assert "Mexico's avocado industry and global trade" in md
        # Author/year line.
        assert "Garcia, M. and Hernandez, L." in md
        assert "2023" in md
        # Excerpt rendered as a blockquote (> prefix).
        assert "> Jalisco's accession" in md

    def test_skips_rag_section_when_no_hits(self):
        result = AnalysisResult(
            parsed=ParsedAnalysis(),
            formatted="", raw="", turn_number=1,
        )
        md = render_markdown(result)
        assert "## Evidence from Literature" not in md

    # ------------------------------------------------------------------
    # PR #32: polish changes.
    # ------------------------------------------------------------------

    def test_title_uses_original_query_when_available(self, minimal_result):
        md = render_markdown(minimal_result)
        first_line = md.splitlines()[0]
        # The fixture sets _original_query_for_export = "Mexico
        # avocado exports to the United States, focusing on the
        # Jalisco production region".
        assert first_line == (
            "# Metacoupling Analysis: Mexico avocado exports to the "
            "United States, focusing on the Jalisco production region"
        )

    def test_title_falls_back_to_topic_when_no_original_query(self):
        """When _original_query_for_export is not attached, the
        title falls back to the older `{focal}: {topic}` heuristic
        so unit tests building AnalysisResult directly still get a
        sensible title."""
        from ._helpers import make_parsed_analysis

        parsed = make_parsed_analysis(
            coupling_classification="Telecoupling case.",
            coupling_type="telecoupling",
        )
        result = AnalysisResult(
            parsed=parsed, formatted="", raw="", turn_number=1,
        )
        md = render_markdown(result)
        first_line = md.splitlines()[0]
        # Fallback shape: "Metacoupling Analysis — Case: ..."
        assert "Metacoupling Analysis" in first_line
        assert "—" in first_line  # the em-dash divider of the fallback

    def test_systems_subfield_labels_are_bold(self, minimal_result):
        md = render_markdown(minimal_result)
        # Bold (`**Label**:`), not italic (`*Label*:`).
        assert "**Human subsystem**:" in md or "**Geographic scope**:" in md
        # Quick negative check: no italic-form sub-labels remain.
        assert "*Geographic scope*:" not in md
        assert "*Description*:" not in md

    def test_agents_level_label_is_bold(self, minimal_result):
        md = render_markdown(minimal_result)
        # The fixture has "Individuals / Households" and "Firms /
        # Traders / Corporations" levels.  Bold not italic.
        assert "**Individuals / Households**" in md
        assert "**Firms / Traders / Corporations**" in md
        # Italic form should NOT be present.
        assert "*Individuals / Households*" not in md.replace(
            "**Individuals / Households**", ""
        )

    def test_flows_grouped_by_category_with_subsections(self, minimal_result):
        md = render_markdown(minimal_result)
        # The fixture flows are matter + capital + information.
        # PR #32 renders these as #### subsections under §4.2 Flows.
        assert "#### 4.2.1 Matter" in md
        assert "#### 4.2.2 Capital" in md
        assert "#### 4.2.3 Information" in md
        # The OLD flat shape ("1. **[Matter]**") should be gone.
        assert "**[Matter]**" not in md

    def test_flows_preserves_canonical_category_order(self):
        """Matter, Capital, Information come in that order even
        when the LLM emits them out of order."""
        from ._helpers import make_parsed_analysis

        parsed = make_parsed_analysis(
            flows=[
                {"category": "information", "direction": "A → B",
                 "description": ""},
                {"category": "capital", "direction": "B → A",
                 "description": ""},
                {"category": "matter", "direction": "A → B",
                 "description": ""},
            ],
            coupling_type="telecoupling",
        )
        parsed.map_data = None  # force fallback to CouplingSection flows
        result = AnalysisResult(
            parsed=parsed, formatted="", raw="", turn_number=1,
        )
        md = render_markdown(result)
        # Find each category heading's position.
        matter_pos = md.find("4.2.1 Matter")
        capital_pos = md.find("4.2.2 Capital")
        information_pos = md.find("4.2.3 Information")
        assert 0 < matter_pos < capital_pos < information_pos

    def test_causes_categories_are_separate_subsections(self, minimal_result):
        md = render_markdown(minimal_result)
        # Fixture causes: "Economic" + "Political / Institutional".
        assert "#### 4.4.1 Economic" in md
        assert "#### 4.4.2 Political / Institutional" in md
        # Each item is its own bullet (not jammed into one paragraph
        # like the old docx behavior).
        assert "- Rising US demand for avocados." in md

    def test_effects_categories_are_separate_subsections(self, minimal_result):
        md = render_markdown(minimal_result)
        assert "#### 4.5.1 Economic" in md
        assert "#### 4.5.2 Environmental" in md
        assert "- Deforestation pressure in highland forests." in md


# ---------------------------------------------------------------------------
# Word (.docx) renderer
# ---------------------------------------------------------------------------


# Skip the whole class when python-docx is unavailable.
docx = pytest.importorskip("docx")


class TestRenderDocx:
    def test_produces_file_at_given_path(self, minimal_result, tmp_path):
        out = tmp_path / "report.docx"
        path = render_docx(minimal_result, path=out)
        assert path == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_default_path_when_none_given(self, minimal_result, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        path = render_docx(minimal_result)
        assert path.name == "metacoupling_report.docx"
        assert path.exists()

    def test_contains_expected_headings(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "report.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        # Collect every heading-style paragraph's text.
        headings = [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        assert any("Abstract" in h for h in headings)
        assert any("1. Coupling Classification" in h for h in headings)
        assert any("4. Telecoupling Analysis" in h for h in headings)

    def test_contains_flows_table(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "report.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        # At least one table should exist for flows.
        assert doc.tables, "no tables in document"
        # Find a table with the expected header row.
        flows_table = None
        for table in doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if (
                "Category" in header_cells
                and "Source" in header_cells
                and "Target" in header_cells
            ):
                flows_table = table
                break
        assert flows_table is not None
        # At least one data row beyond the header.
        assert len(flows_table.rows) >= 2

    def test_contains_web_sources_table(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "report.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        # Look for a table whose header includes ID + URL.
        found = False
        for table in doc.tables:
            header_cells = [cell.text for cell in table.rows[0].cells]
            if "ID" in header_cells and "URL" in header_cells:
                found = True
                break
        assert found

    def test_to_docx_method_delegates(self, minimal_result, tmp_path):
        out = tmp_path / "method_call.docx"
        path = minimal_result.to_docx(out)
        assert path == out
        assert out.exists()

    # PR #31 follow-up: RAG evidence rendering in docx.

    def test_contains_rag_evidence_headings(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "rag.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        headings = [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ]
        # Section heading + per-hit subheading.
        assert any("Evidence from Literature" in h for h in headings)
        assert any("[T1:1]" in h for h in headings)
        assert any("Mexico's avocado industry" in h for h in headings)

    def test_rag_excerpt_text_present_in_docx(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "rag_text.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jalisco's accession" in all_text
        # Author/year line rendered.
        assert "Garcia, M. and Hernandez, L." in all_text

    # ------------------------------------------------------------------
    # PR #32: polish changes in docx.
    # ------------------------------------------------------------------

    def test_title_uses_original_query(self, minimal_result, tmp_path):
        from docx import Document

        out = tmp_path / "polish_title.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        # Heading 0 (Title) paragraph should contain the user's query.
        title_paragraphs = [
            p for p in doc.paragraphs if p.style.name == "Title"
        ]
        assert title_paragraphs, "no Title-styled paragraph found"
        title_text = title_paragraphs[0].text
        assert "Metacoupling Analysis" in title_text
        assert "Mexico avocado exports" in title_text

    def test_systems_subfield_labels_use_bold_runs(
        self, minimal_result, tmp_path,
    ):
        from docx import Document

        out = tmp_path / "polish_systems.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        # Find a sub-field paragraph (style "List Bullet") whose first
        # run is bold (e.g., "Geographic scope: ").
        found_bold_label = False
        for p in doc.paragraphs:
            if p.style.name != "List Bullet" or not p.runs:
                continue
            first_run = p.runs[0]
            if (
                first_run.bold
                and first_run.text.rstrip().endswith(":")
                and any(label in first_run.text for label in (
                    "Human subsystem", "Natural subsystem",
                    "Geographic scope", "Description",
                ))
            ):
                found_bold_label = True
                break
        assert found_bold_label, (
            "expected a bold-prefix sub-field label in §N.1 Systems"
        )

    def test_flows_have_per_category_h3_headings(
        self, minimal_result, tmp_path,
    ):
        from docx import Document

        out = tmp_path / "polish_flows.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        h3_texts = [
            p.text for p in doc.paragraphs
            if p.style.name == "Heading 3"
        ]
        # Fixture flows have matter + capital + information, rendered
        # in §4 (Telecoupling).
        assert any("4.2.1 Matter" in t for t in h3_texts)
        assert any("4.2.2 Capital" in t for t in h3_texts)
        assert any("4.2.3 Information" in t for t in h3_texts)

    def test_causes_have_per_category_h3_headings(
        self, minimal_result, tmp_path,
    ):
        from docx import Document

        out = tmp_path / "polish_causes.docx"
        render_docx(minimal_result, path=out)
        doc = Document(str(out))
        h3_texts = [
            p.text for p in doc.paragraphs
            if p.style.name == "Heading 3"
        ]
        # Causes: Economic, Political / Institutional.
        assert any("4.4.1 Economic" in t for t in h3_texts)
        assert any(
            "4.4.2 Political / Institutional" in t for t in h3_texts
        )
        # Each item is its own List Bullet paragraph (not jammed
        # into one semicolon paragraph as before).
        bullet_texts = [
            p.text for p in doc.paragraphs
            if p.style.name == "List Bullet"
        ]
        assert any(
            "Rising US demand for avocados." in t for t in bullet_texts
        )


# ---------------------------------------------------------------------------
# PR #33: defensive helpers for fragmented parser output.
# ---------------------------------------------------------------------------


class TestMergeFragmentedFlows:
    """``_merge_fragmented_flows`` stitches the parser's 3-entry
    header / direction / description pattern back into single flow
    dicts.  Modelled on the real Mexico avocado trace at
    runs/pr31_export_2026-05-22_105728/formatted.txt lines 23-52."""

    def test_merges_3entry_pattern_into_single_flow(self):
        # Pattern observed in the live trace:
        #   Header (canonical category, direction="Unspecified",
        #           description="Matter flow")
        #   Direction-only (category="unspecified", direction="X→Y",
        #                   description="")
        #   Description-only (category="unspecified",
        #                     direction="Unspecified",
        #                     description="actual prose...")
        fragmented = [
            {"category": "matter", "direction": "Unspecified",
             "description": "Matter flow"},
            {"category": "unspecified",
             "direction": "Orchards → Packinghouses (within Jalisco)",
             "description": ""},
            {"category": "unspecified", "direction": "Unspecified",
             "description": "Harvested Hass avocados moved locally."},
        ]
        merged = _merge_fragmented_flows(fragmented)
        assert len(merged) == 1
        assert merged[0]["category"] == "matter"
        assert merged[0]["direction"] == (
            "Orchards → Packinghouses (within Jalisco)"
        )
        assert "Harvested Hass avocados" in merged[0]["description"]

    def test_merges_multiple_categories_independently(self):
        """Matter and Capital each have their own 3-entry block;
        merge produces 2 final flows in canonical order."""
        fragmented = [
            {"category": "matter", "direction": "Unspecified",
             "description": "Matter flow"},
            {"category": "unspecified", "direction": "A → B",
             "description": ""},
            {"category": "capital", "direction": "Unspecified",
             "description": "Capital flow"},
            {"category": "unspecified", "direction": "B → A",
             "description": "Payments"},
        ]
        merged = _merge_fragmented_flows(fragmented)
        assert len(merged) == 2
        assert merged[0]["category"] == "matter"
        assert merged[0]["direction"] == "A → B"
        assert merged[1]["category"] == "capital"
        assert merged[1]["direction"] == "B → A"
        assert merged[1]["description"] == "Payments"

    def test_idempotent_on_clean_flows(self):
        """Clean flows (canonical category + real direction +
        non-placeholder description) pass through unchanged."""
        clean = [
            {"category": "matter", "direction": "Mexico → USA",
             "description": "Fresh avocados", "source": "MEX",
             "target": "USA", "bidirectional": False},
            {"category": "capital", "direction": "USA → Mexico",
             "description": "Payments", "source": "USA",
             "target": "MEX", "bidirectional": False},
        ]
        merged = _merge_fragmented_flows(clean)
        assert merged == clean

    def test_empty_input_returns_empty(self):
        assert _merge_fragmented_flows([]) == []

    def test_orphan_unspecified_without_header_passes_through(self):
        """A sub-flow without a preceding canonical header still
        gets emitted (otherwise we'd silently drop data)."""
        orphan = [
            {"category": "unspecified",
             "direction": "Orphan direction",
             "description": "no header above me"},
        ]
        merged = _merge_fragmented_flows(orphan)
        assert merged == orphan


class TestSplitCollapsedCausesEffects:
    """``_split_collapsed_causes_effects`` detects the parser's
    everything-under-General pattern and splits by inline Liu
    framework category names."""

    def test_splits_general_into_real_categories(self):
        collapsed = {
            "General": [
                "Economic",
                "Strong U.S. demand and price incentives.",
                "Anticipated revenue from market access.",
                "Political / Institutional",
                "Phytosanitary requirements from SENASICA.",
                "Hydrological",
                "Local water availability conditioning yields.",
            ]
        }
        split = _split_collapsed_causes_effects(collapsed)
        # Three real categories, not the bogus "General" key.
        assert "General" not in split
        assert set(split.keys()) == {
            "Economic", "Political / Institutional", "Hydrological",
        }
        assert len(split["Economic"]) == 2
        assert split["Hydrological"] == [
            "Local water availability conditioning yields."
        ]

    def test_idempotent_on_multi_key_input(self):
        """A dict already keyed by real categories passes through
        unchanged (parser worked correctly)."""
        clean = {
            "Economic": ["Strong demand."],
            "Political / Institutional": ["Regulatory protocol."],
        }
        assert _split_collapsed_causes_effects(clean) == clean

    def test_safety_belt_skips_when_under_two_categories(self):
        """If items contain <2 known category names, leave the
        dict alone (heuristic mis-fire protection)."""
        only_one_cat = {
            "General": [
                "Economic",  # just one category-shaped item
                "Strong U.S. demand and price incentives.",
                "Some other item without category shape.",
            ]
        }
        assert _split_collapsed_causes_effects(only_one_cat) == only_one_cat

    def test_empty_input_returns_unchanged(self):
        assert _split_collapsed_causes_effects({}) == {}


# ---------------------------------------------------------------------------
# PR #33: end-to-end Markdown rendering against a broken-parser fixture.
# ---------------------------------------------------------------------------


@pytest.fixture
def broken_parser_result() -> AnalysisResult:
    """A fixture matching the actual Mexico avocado trace's parser
    output: fragmented flows + flattened General-keyed causes.
    Verifies that the new helpers turn the broken shape into a
    clean rendering."""
    parsed = make_parsed_analysis(
        coupling_classification=(
            "- Intracoupling — within Jalisco's coupled human-natural system.\n"
            "- Pericoupling — Jalisco is adjacent to Michoacán.\n"
            "- Telecoupling — Exports from Jalisco to the United States."
        ),
        systems={
            "focal": {
                "name": "Avocado production region of Jalisco, Mexico",
                "human_subsystem": "Smallholder growers; packinghouses.",
                "natural_subsystem": "Pine-oak ecosystems; soils.",
                "geographic_scope": "Avocado-producing municipalities.",
            },
        },
        flows=[
            # Fragmented Matter block (3 entries)
            {"category": "matter", "direction": "Unspecified",
             "description": "Matter flow"},
            {"category": "unspecified",
             "direction": "Orchards → Packinghouses (within Jalisco)",
             "description": ""},
            {"category": "unspecified", "direction": "Unspecified",
             "description": "Harvested Hass avocados moved locally [T1:W8]."},
            # Fragmented Capital block (3 entries)
            {"category": "capital", "direction": "Unspecified",
             "description": "Capital flow"},
            {"category": "unspecified",
             "direction": "Exporter financing → Orchards",
             "description": ""},
            {"category": "unspecified", "direction": "Unspecified",
             "description": "Investments in compliance [T1:W3]."},
        ],
        causes={
            # Flattened parser output: all categories + items under "General"
            "General": [
                "Economic",
                "Strong U.S. demand and price incentives [T1:W7].",
                "Anticipated revenue from market access [T1:W3].",
                "Political / Institutional",
                "Phytosanitary requirements from SENASICA [T1:W8].",
                "Hydrological",
                "Local water availability conditioning yields [T1:W2].",
            ],
        },
        coupling_type="intracoupling",
    )
    return AnalysisResult(
        parsed=parsed, formatted="", raw="", turn_number=1,
    )


class TestBrokenParserResultRendering:
    """End-to-end: render the broken-parser fixture and verify the
    output looks clean (the symptoms the user reported are gone)."""

    def test_markdown_flows_have_real_direction_and_description(
        self, broken_parser_result,
    ):
        md = render_markdown(broken_parser_result)
        # §2.2.1 Matter has the merged direction + description.
        assert "#### 2.2.1 Matter" in md
        assert "Orchards → Packinghouses" in md
        assert "Harvested Hass avocados" in md
        # §2.2.2 Capital likewise.
        assert "#### 2.2.2 Capital" in md
        assert "Exporter financing → Orchards" in md
        assert "Investments in compliance" in md

    def test_markdown_no_placeholder_only_flow_items(
        self, broken_parser_result,
    ):
        """The 'Matter flow' / 'Capital flow' placeholder items
        from the parser must NOT appear as standalone numbered
        list items in the rendered output."""
        md = render_markdown(broken_parser_result)
        # The merged flow uses its real direction + description,
        # not the placeholder description.
        for placeholder in (
            "1. Matter flow", "1. Capital flow",
            "1. Information flow", "1. Energy flow",
            "1. People flow", "1. Organism flow",
        ):
            assert placeholder not in md, (
                f"placeholder {placeholder!r} should not appear in output"
            )

    def test_markdown_no_unspecified_mega_list(
        self, broken_parser_result,
    ):
        """The §2.2.7 Unspecified mega-list should not appear --
        all flows merged into their canonical categories."""
        md = render_markdown(broken_parser_result)
        assert "#### 2.2.7 Unspecified" not in md

    def test_markdown_causes_split_into_real_categories(
        self, broken_parser_result,
    ):
        md = render_markdown(broken_parser_result)
        # Real categories appear as §2.4.K subsections.
        assert "#### 2.4.1 Economic" in md
        assert "#### 2.4.2 Political / Institutional" in md
        assert "#### 2.4.3 Hydrological" in md
        # The bogus "General" key is gone.
        assert "#### 2.4.1 General" not in md


class TestDocxClassificationBullets:
    """PR #33: §1 Coupling Classification with `- Intra / Peri / Tele`
    bullets renders as 3 List Bullet paragraphs in docx (not one
    giant Normal paragraph)."""

    def test_classification_bullets_render_as_list_bullets(
        self, broken_parser_result, tmp_path,
    ):
        from docx import Document

        out = tmp_path / "cls_bullets.docx"
        render_docx(broken_parser_result, path=out)
        doc = Document(str(out))
        # After §1 heading, expect 3 List Bullet paragraphs (one per
        # coupling type) -- not a single Normal block.
        bullet_texts = [
            p.text for p in doc.paragraphs
            if p.style.name == "List Bullet"
        ]
        # The 3 classification bullets should be among them.
        assert any(
            t.startswith("Intracoupling") for t in bullet_texts
        )
        assert any(
            t.startswith("Pericoupling") for t in bullet_texts
        )
        assert any(
            t.startswith("Telecoupling") for t in bullet_texts
        )

    def test_classification_single_paragraph_falls_back(self, tmp_path):
        """When classification is a single non-bulleted string, the
        docx still renders it as a regular paragraph (no spurious
        bullet split)."""
        from docx import Document

        parsed = make_parsed_analysis(
            coupling_classification="A single classification sentence.",
            coupling_type="telecoupling",
        )
        result = AnalysisResult(
            parsed=parsed, formatted="", raw="", turn_number=1,
        )
        out = tmp_path / "cls_single.docx"
        render_docx(result, path=out)
        doc = Document(str(out))
        # Look for the classification text in any Normal paragraph.
        normal_texts = [
            p.text for p in doc.paragraphs
            if p.style.name == "Normal"
        ]
        assert any(
            "A single classification sentence." in t
            for t in normal_texts
        )
